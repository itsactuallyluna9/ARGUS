from io import BytesIO, StringIO
import mimetypes
from time import time

from bs4 import BeautifulSoup
import ollama
import requests
import trafilatura
from trafilatura.readability_lxml import is_probably_readerable
from playwright.async_api import Error, async_playwright
from playwright_stealth import Stealth
from pdf_oxide import PdfDocument
from tempfile import NamedTemporaryFile
import pyexcel
import docx


class ScraperError(Exception):
    pass


class UnsupportedContentTypeError(ScraperError):
    pass


def ttl_cache(ttl_seconds):
    # TODO: async support
    def decorator(func):
        cache = {}

        def wrapper(*args, **kwargs):
            cache_key = (func.__name__,) + tuple(args) + tuple(kwargs.items())
            now = time()
            cached_result = cache.get(cache_key)
            if cached_result is None or now - cached_result["timestamp"] > ttl_seconds:
                cached_result = func(*args, **kwargs)
                cache[cache_key] = {"result": cached_result, "timestamp": now}
                # cleanup old cache entries
                for key in list(cache.keys()):
                    if now - cache[key]["timestamp"] > ttl_seconds:
                        del cache[key]
                return cached_result
            return cached_result["result"]

        wrapper.clear_cache = lambda: cache.clear()
        return wrapper

    return decorator


async def get_page(url: str) -> tuple[dict[str, str], str]:
    import logging

    logger = logging.getLogger(__name__)
    metadata = {}
    content_type, content, used_fallback = await get_source(url)
    metadata["content_type"] = content_type
    metadata["content_length"] = len(content)
    metadata["used_fallback"] = used_fallback

    logger.info(f"Fetched content from {url} with content type {content_type} and length {len(content)} (used fallback: {used_fallback})")

    match content_type:
        case "text/html" | "application/xhtml+xml":
            html = content.decode("utf-8", errors="ignore")
            metadata.update(extract_html_metadata(html))
            if is_probably_readerable(html):
                metadata["readerable"] = True
                metadata.update(trafilatura.extract_metadata(html).as_dict())
                del metadata["body"]
                del metadata["comments"]
                del metadata["commentsbody"]
                for key in set(metadata.keys()):
                    if metadata[key] is None:
                        del metadata[key]
                cleaned = trafilatura.extract(html, output_format="markdown")
                if cleaned:
                    return metadata, cleaned
            # maybe we got something trafilatura doesn't like, but that still has content?
            metadata["readerable"] = False
            # well. readerlm?
            cleaned = ollama.generate("reader-lm", html).response
            return metadata, cleaned
        case "application/pdf":
            return metadata, extract_pdf(content)
        case "application/msword":
            pass  # TODO: handle classic word docs
        case "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return metadata, extract_docx(content)
        case "text/csv" | "application/vnd.ms-excel" | "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet" | "application/vnd.oasis.opendocument.spreadsheet":
            return metadata, extract_sheet(content, content_type)
        case "application/vnd.ms-powerpoint" | "application/vnd.openxmlformats-officedocument.presentationml.presentation":
            pass  # TODO: handle powerpoint presentations
        case "application/vnd.oasis.opendocument.text":
            return metadata, extract_openword(content)
        case "application/vnd.oasis.opendocument.presentation":
            pass  # TODO: handle open document presentation
        case "text/plain" | "text/markdown":
            return metadata, content.decode()  # fair enough
        case "text/rtf":
            pass  # TODO: handle rtf
        case "application/json":
            return metadata, content.decode()  # sure, why not
        # images?
        case _:
            raise UnsupportedContentTypeError(f"Unsupported content type: {content_type}")
    raise ScraperError(f"Failed to extract content from page with content type: {content_type}")


# @ttl_cache(5 * 60)
async def get_source(url: str) -> tuple[str, bytes, bool]:
    try:
        # TODO: swap this out with async!
        resp = requests.get(url, timeout=60)
    except requests.RequestException:
        return *await get_source_chrome(url), True

    content_type = resp.headers.get("Content-Type", "").lower().split(";")[0].strip()

    if not resp.ok:
        return *await get_source_chrome(url), True

    return content_type, resp.content, False


async def get_source_chrome(url: str) -> tuple[str, bytes]:
    # this should beat the "the simplest of bot detection methods."
    # it does work on cornell.
    # this is probably a commentary on something, but i'm not sure what.
    async with Stealth().use_async(async_playwright()) as p:
        browser = await p.chromium.launch(headless=True)
        page = await browser.new_page()

        downloads = []
        page.on("download", lambda download: downloads.append(download))

        response = None
        try:
            # okay. here we have a couple of possibilities.
            # 1) this is a normal page.
            # 2) this is actually a pdf or something, and chrome being chrome will do its in-browser pdf rendering thing.
            # 3) we'll get a download
            # 4) something will go wrong, and we can load the page, but its worthless
            # 5) something will go wrong, and we cant actually load the page

            # do we get a download?
            try:
                response = await page.goto(url)
            except Error as e:
                # alright: we might have something?
                if "Download is starting" in str(e):
                    # yay! the browser signalled a download is starting.
                    # wait briefly for the download event so we can capture it
                    try:
                        download = await page.wait_for_event("download", timeout=10000)
                        downloads.append(download)
                    except Exception:
                        # if waiting fails, fall through and let the later
                        # `if downloads:` check handle any downloads that arrived.
                        pass
                else:
                    raise
            if downloads:
                # yay! we got something (after page load... what??)
                download = downloads[0]
                download_path = download.path()
                if not download_path:
                    raise ScraperError("Browser download completed but file path was not available.")

                with open(download_path, "rb") as f:
                    content = f.read()

                content_type = ""
                if response:
                    content_type = response.headers.get("content-type", ";").lower().split(";")[0].strip()
                content_type = content_type or mimetypes.guess_type(download.suggested_filename)[0] or "application/octet-stream"
                return content_type, content

            # get the content type from response headers (handles redirects automatically via page.goto)
            content_type = "text/html"
            if response:
                content_type = response.headers.get("content-type", "text/html").lower().split(";")[0].strip()

            # always get bytes from rendered page content
            content = (await page.content()).encode("utf-8")

        except Exception as e:
            # raise ScraperError(f"Failed to fetch source with headless browser: {e}")
            return "text/plain", f"Failed to fetch source with headless browser: {e}".encode()
        
        finally:
            await browser.close()

    return content_type, content


# MARK: - Parsers


def extract_html_metadata(html: str) -> dict[str, str]:
    soup = BeautifulSoup(html, "html.parser")
    metadata = {}
    # title and description
    title_tag = soup.find("title")
    if title_tag:
        metadata["title"] = title_tag.text.strip()
    description_tag = soup.find("meta", attrs={"name": "description"})
    if description_tag and description_tag.get("content"):
        metadata["description"] = description_tag["content"].strip()
    # meta
    # opengraph
    for meta in soup.find_all("meta"):
        # og:
        if meta.get("property", "").startswith("og:") and meta.get("content"):
            key = meta["property"][3:]  # remove "og:" prefix
            metadata[key] = meta["content"].strip()
        # twitter:
        elif meta.get("name", "").startswith("twitter:") and meta.get("content"):
            key = meta["name"][8:]  # remove "twitter:" prefix
            if key in metadata:
                key = "twitter_" + key  # avoid collisions with og: tags
            metadata[key] = meta["content"].strip()
    return metadata


def extract_pdf(content: bytes) -> str:
    with NamedTemporaryFile(suffix=".pdf") as tmpfile:
        tmpfile.write(content)
        tmpfile.flush()
        tmpfile.seek(0)
        with PdfDocument(tmpfile.name) as doc:
            # TODO: look into OCR support
            # cant count on everyone having nice pdfs :c
            return doc.to_markdown_all(detect_headings=True, embed_images=False)


def extract_sheet(content: bytes, content_type: str) -> str:
    lookup = {
        "application/vnd.ms-excel": "xls",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.oasis.opendocument.spreadsheet": "ods",
        "text/csv": "csv",
    }
    extension = lookup.get(content_type)
    if extension is None:
        raise NotImplementedError(f"scraper - unsupported spreadsheet content type: {content_type} (how did we get here?)")
    with NamedTemporaryFile(suffix=f".{extension}") as tmpfile:
        tmpfile.write(content)
        tmpfile.flush()
        tmpfile.seek(0)
        book_dict = pyexcel.get_book_dict(file_name=tmpfile.name)
        markdown = "*A table was found. Here is the content of the table in markdown format:*\n\n"
        for sheet_name, sheet_contents in book_dict.items():
            # sheet contents: list of lists. [0] is row 1, [0][0] is cell A1
            markdown += f"## {sheet_name}\n\n"
            # now we get to generate a table by hand. fun.
            if len(sheet_contents) == 0:
                markdown += "*This sheet is empty.*\n\n"
                continue
            # alright, we can't really assume there's a formal header row.
            # we're... just going to not worry about it! it cant harm us
            # if we don't think about it.
            for row in sheet_contents:
                markdown += "| " + " | ".join(str(cell) for cell in row) + " |\n"
            markdown += "\n\n"
        return markdown


def extract_docx(content: bytes) -> str:
    with BytesIO(content) as tmpfile, StringIO() as output:
        doc = docx.Document(tmpfile)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                # we're gonna try to keep the formatting here!
                # might be useful?
                # TODO: consider more formatting options!
                # namely: lists.
                wrapping = ""
                if run.bold or (paragraph.style and paragraph.style.font.bold):
                    wrapping += "**"
                if run.italic or (paragraph.style and paragraph.style.font.italic):
                    wrapping += "*"
                if run.underline or (paragraph.style and paragraph.style.font.underline):
                    wrapping += "__"
                output.write(f"{wrapping}{run.text}{wrapping}")
            output.write("\n\n")
        return output.getvalue()


if __name__ == "__main__":
    from sys import argv
    import asyncio

    print(asyncio.run(get_page(argv[1]))[1])
